from pathlib import Path
import pdfplumber
from docx import Document


def read_document_text(path: Path) -> str:
    """Extract plain text from a PDF or DOCX file."""
    if not path.exists():
        raise FileNotFoundError(path)

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported format: {suffix}")


def _read_pdf(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            if txt:
                parts.append(txt)
    return "\n".join(parts)


def _read_docx(path: Path) -> str:
    """Extract text from body paragraphs, tables (incl. nested), and section headers/footers."""
    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]

    for tbl in doc.tables:
        parts.extend(_walk_table(tbl))

    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                if p.text:
                    parts.append(p.text)

    return "\n".join(parts)


def _walk_table(tbl) -> list[str]:
    """Recursively read text from a docx table including nested tables."""
    out: list[str] = []
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if p.text:
                    out.append(p.text)
            for nested in cell.tables:
                out.extend(_walk_table(nested))
    return out
